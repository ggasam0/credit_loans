#!/usr/bin/env python
"""
Generate an IEEE-style DOCX report from a Markdown source file.

Usage:
  python 6800paper/scripts/md_to_ieee_docx.py
"""

from __future__ import annotations

import argparse
import re
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import List

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


NORMALIZE_URI_REPLACEMENTS = [
    (
        "http://purl.oclc.org/ooxml/wordprocessingml/main",
        "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    ),
    (
        "http://purl.oclc.org/ooxml/drawingml/main",
        "http://schemas.openxmlformats.org/drawingml/2006/main",
    ),
    (
        "http://purl.oclc.org/ooxml/drawingml/wordprocessingDrawing",
        "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    ),
    (
        "http://purl.oclc.org/ooxml/officeDocument/math",
        "http://schemas.openxmlformats.org/officeDocument/2006/math",
    ),
    (
        "http://purl.oclc.org/ooxml/officeDocument/relationships/extendedProperties",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties",
    ),
    (
        "http://purl.oclc.org/ooxml/officeDocument/extendedProperties",
        "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties",
    ),
    (
        "http://purl.oclc.org/ooxml/officeDocument/relationships",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    ),
    (
        "http://purl.oclc.org/ooxml/officeDocument/customXml",
        "http://schemas.openxmlformats.org/officeDocument/2006/customXml",
    ),
    (
        "http://purl.oclc.org/ooxml/officeDocument/docPropsVTypes",
        "http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes",
    ),
]


@dataclass
class Block:
    kind: str
    text: str = ""
    rows: List[List[str]] = field(default_factory=list)


@dataclass
class ParsedReport:
    title: str
    author: str
    course: str
    date: str
    abstract: str
    keywords: str
    sections: List[tuple[str, List[Block]]]
    references: List[str]


def normalize_template(template_path: Path, normalized_path: Path) -> Path:
    normalized_path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(template_path, "r") as zin, zipfile.ZipFile(
        normalized_path, "w", zipfile.ZIP_DEFLATED
    ) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.endswith(".xml") or item.filename.endswith(".rels"):
                try:
                    text = data.decode("utf-8")
                    for old, new in NORMALIZE_URI_REPLACEMENTS:
                        text = text.replace(old, new)
                    data = text.encode("utf-8")
                except UnicodeDecodeError:
                    pass
            zout.writestr(item, data)
    return normalized_path


def clean_md_text(line: str) -> str:
    text = line.strip()
    text = text.replace("`", "")
    text = re.sub(r"\s+", " ", text)
    return text


def normalize_heading_text(text: str, level: int) -> str:
    """
    Remove manual heading indices from markdown text because IEEE heading styles
    in the template already auto-number headings.
    """
    t = clean_md_text(text)
    if level == 1:
        # e.g., "I. Introduction", "V. Methods"
        t = re.sub(r"^[IVXLCDM]+\.\s*", "", t, flags=re.IGNORECASE)
    elif level == 2:
        # e.g., "A. 数据来源与清洗", "B. Methods"
        t = re.sub(r"^[A-Z]\.\s*", "", t)
    return t.strip()


def parse_table_rows(table_lines: List[str]) -> List[List[str]]:
    rows: List[List[str]] = []
    for raw in table_lines:
        stripped = raw.strip()
        if not stripped.startswith("|"):
            continue
        cells = [c.strip() for c in stripped.strip("|").split("|")]
        if not cells:
            continue
        if all(re.fullmatch(r":?-{2,}:?", c.replace(" ", "")) for c in cells):
            continue
        rows.append(cells)
    return rows


def parse_blocks(lines: List[str]) -> List[Block]:
    blocks: List[Block] = []
    paragraph_parts: List[str] = []

    def flush_paragraph() -> None:
        if paragraph_parts:
            text = clean_md_text(" ".join(paragraph_parts))
            if text:
                blocks.append(Block(kind="paragraph", text=text))
            paragraph_parts.clear()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            i += 1
            continue

        if stripped.startswith("### "):
            flush_paragraph()
            blocks.append(Block(kind="heading2", text=clean_md_text(stripped[4:])))
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            flush_paragraph()
            blocks.append(Block(kind="numbered", text=clean_md_text(stripped)))
            i += 1
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            blocks.append(Block(kind="bullet", text=clean_md_text(stripped[2:])))
            i += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph()
            table_lines: List[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table_rows(table_lines)
            if rows:
                blocks.append(Block(kind="table", rows=rows))
            continue

        paragraph_parts.append(stripped)
        i += 1

    flush_paragraph()
    return blocks


def parse_references(ref_lines: List[str]) -> List[str]:
    refs: List[str] = []
    current = ""
    for raw in ref_lines:
        line = raw.strip()
        if not line:
            continue
        if re.match(r"^\[\d+\]\s+", line):
            if current:
                refs.append(clean_md_text(current))
            current = line
        else:
            if current:
                current += " " + line
    if current:
        refs.append(clean_md_text(current))
    return refs


def parse_markdown(md_path: Path) -> ParsedReport:
    lines = md_path.read_text(encoding="utf-8").splitlines()

    title = ""
    author = ""
    course = ""
    date = ""
    abstract = ""
    keywords = ""
    sections: List[tuple[str, List[Block]]] = []
    references: List[str] = []

    for line in lines:
        stripped = line.strip().lstrip("\ufeff")
        if stripped.startswith("# ") and not title:
            title = clean_md_text(stripped[2:])
        elif stripped.startswith("作者：") and not author:
            author = clean_md_text(stripped)
        elif stripped.startswith("课程：") and not course:
            course = clean_md_text(stripped)
        elif stripped.startswith("日期：") and not date:
            date = clean_md_text(stripped)

    heading_indices = []
    for idx, line in enumerate(lines):
        if line.startswith("## "):
            heading_indices.append((idx, line[3:].strip()))
    heading_indices.append((len(lines), "__END__"))

    # Abstract and keywords
    abstract_start = next(
        (i for i, h in heading_indices if h in {"Abstract", "\u6458\u8981"}), None
    )
    if abstract_start is not None:
        abstract_end = next(
            (i for i, h in heading_indices if i > abstract_start), len(lines)
        )
        abstract_lines = []
        for line in lines[abstract_start + 1 : abstract_end]:
            stripped = line.strip()
            if stripped.startswith("**Keywords**:") or stripped.startswith(
                "**\u5173\u952e\u8bcd**:"
            ):
                keywords = clean_md_text(stripped.split(":", 1)[1])
            elif stripped:
                abstract_lines.append(stripped)
        abstract = clean_md_text(" ".join(abstract_lines))

    # Sections and references
    for i in range(len(heading_indices) - 1):
        start_idx, heading = heading_indices[i]
        end_idx, _ = heading_indices[i + 1]
        body_lines = lines[start_idx + 1 : end_idx]

        if heading in {"Abstract", "\u6458\u8981"}:
            continue
        if heading in {"References", "\u53c2\u8003\u6587\u732e"}:
            references = parse_references(body_lines)
            continue
        if heading.startswith("Appendix"):
            continue
        blocks = parse_blocks(body_lines)
        sections.append((heading, blocks))

    return ParsedReport(
        title=title,
        author=author,
        course=course,
        date=date,
        abstract=abstract,
        keywords=keywords,
        sections=sections,
        references=references,
    )


def delete_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)


def set_two_column_sections(doc: Document) -> None:
    for idx, section in enumerate(doc.sections):
        if idx == 0:
            # Keep first section layout from template (title area).
            continue
        sect_pr = section._sectPr
        cols = sect_pr.xpath("./w:cols")
        if cols:
            col = cols[0]
        else:
            col = OxmlElement("w:cols")
            sect_pr.append(col)
        col.set(qn("w:num"), "2")
        col.set(qn("w:space"), "720")  # 0.5 inch


def fill_docx_from_report(report: ParsedReport, template_path: Path, output_path: Path) -> None:
    doc = Document(str(template_path))

    # Title
    if doc.paragraphs:
        doc.paragraphs[0].text = report.title or "6800 Final Report"
        doc.paragraphs[0].style = "paper title"

    # Author block
    author_lines = [x for x in [report.author, report.course, report.date] if x]
    author_text = (
        "\n".join(author_lines)
        if author_lines
        else "\u4f5c\u8005\uff1a<\u59d3\u540d/\u5b66\u53f7>"
    )

    author_target = None
    for p in doc.paragraphs:
        if p.style.name == "Author" and "line 1:" in p.text:
            author_target = p
            break
    if author_target is None:
        for p in doc.paragraphs:
            if p.style.name == "Author":
                author_target = p
                break
    if author_target is not None:
        author_target.text = author_text

    # Remove note and extra placeholder author blocks
    for p in doc.paragraphs:
        if p.style.name == "Author" and p.text.strip().startswith("*Note:"):
            p.text = ""
    for p in list(doc.paragraphs):
        if (
            p.style.name == "Author"
            and p is not author_target
            and "line 1:" in p.text
        ):
            delete_paragraph(p)

    # Abstract/keywords in template
    abstract_idx = next(
        (i for i, p in enumerate(doc.paragraphs) if p.style.name == "Abstract"), None
    )
    keywords_idx = next(
        (i for i, p in enumerate(doc.paragraphs) if p.style.name == "Keywords"), None
    )
    if abstract_idx is None or keywords_idx is None:
        raise RuntimeError("Template missing Abstract/Keywords styles.")

    doc.paragraphs[abstract_idx].text = (
        f"\u6458\u8981\u2014{report.abstract}"
        if report.abstract
        else "\u6458\u8981\u2014"
    )
    doc.paragraphs[keywords_idx].text = (
        f"\u5173\u952e\u8bcd\u2014{report.keywords}"
        if report.keywords
        else "\u5173\u952e\u8bcd\u2014"
    )

    # Remove all content after keywords (template demo text)
    for p in doc.paragraphs[keywords_idx + 1 :][::-1]:
        delete_paragraph(p)

    # Remove sample tables from template
    for tbl in doc.tables:
        t = tbl._element
        t.getparent().remove(t)

    # Add body sections
    for section_title, blocks in report.sections:
        doc.add_paragraph(normalize_heading_text(section_title, 1), style="Heading 1")
        for block in blocks:
            if block.kind == "heading2":
                doc.add_paragraph(normalize_heading_text(block.text, 2), style="Heading 2")
            elif block.kind == "paragraph":
                doc.add_paragraph(block.text, style="Body Text")
            elif block.kind == "bullet":
                doc.add_paragraph(block.text, style="bullet list")
            elif block.kind == "numbered":
                doc.add_paragraph(block.text, style="Body Text")
            elif block.kind == "table":
                rows = block.rows
                if not rows:
                    continue
                max_cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=max_cols)
                for r_idx, row in enumerate(rows):
                    for c_idx in range(max_cols):
                        text = row[c_idx] if c_idx < len(row) else ""
                        table.cell(r_idx, c_idx).text = text

    # References
    if report.references:
        doc.add_paragraph("References", style="Heading 5")
        for ref in report.references:
            doc.add_paragraph(ref, style="references")

    set_two_column_sections(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate IEEE DOCX from Markdown.")
    parser.add_argument(
        "--md",
        default="6800paper/6800main.md",
        help="Markdown input path.",
    )
    parser.add_argument(
        "--template",
        default="6800paper/conference-template-letter.docx",
        help="IEEE template DOCX path.",
    )
    parser.add_argument(
        "--out",
        default="6800paper/6800_final_report_ieee.docx",
        help="Output DOCX path.",
    )
    parser.add_argument(
        "--tmp-normalized",
        default="tmp/docs/conference-template-letter.normalized.docx",
        help="Temporary normalized template path.",
    )
    args = parser.parse_args()

    md_path = Path(args.md)
    template_path = Path(args.template)
    out_path = Path(args.out)
    normalized_path = Path(args.tmp_normalized)

    report = parse_markdown(md_path)
    normalized_template = normalize_template(template_path, normalized_path)
    fill_docx_from_report(report, normalized_template, out_path)

    print(f"Generated: {out_path}")


if __name__ == "__main__":
    main()
